from fastapi import APIRouter, HTTPException, Depends, Header, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
import io
import os
import json
import re
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from openai import AsyncOpenAI
from auth import get_current_user

router = APIRouter()

# ══════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════

def _openai_client(api_key: str | None = None) -> AsyncOpenAI:
    key = api_key or os.getenv("OPENAI_API_KEY")
    if not key:
        raise RuntimeError("OpenAI API key non configurata.")
    return AsyncOpenAI(api_key=key)


def _slugify(name: str) -> str:
    """Converte nome pagina in slug URL."""
    # Gestisce il separatore "–" o "-" per sotto-pagine
    parts = re.split(r"\s*[–-]\s*", name, maxsplit=1)
    segments = []
    for part in parts:
        slug = part.strip().lower()
        slug = re.sub(r"[^a-z0-9\s]", "", slug)
        slug = re.sub(r"\s+", "-", slug).strip("-")
        segments.append(slug)
    if len(segments) == 2:
        return "/" + segments[0] + "/" + segments[1]
    if segments[0] in ("home", "homepage"):
        return "/"
    return "/" + (segments[0] if segments else "")


def _parse_sections(doc: Document) -> list[dict]:
    """Estrae sezioni pagina dal documento .docx."""
    pages: list[dict] = []
    current_page: str | None = None
    current_url: str = ""
    current_lines: list[str] = []

    # Parole da ignorare come nome pagina
    SKIP_NAMES = {
        "INVIA", "CONTATTI FORM",
        "TITLE", "TITLE:", "DESCRIPTION", "DESCRIPTION:",
        "H1", "H1:", "H2", "H2:", "H3", "H3:",
        "SLOGAN", "SLOGAN:",
        "URL", "URL:",
        "NOTE", "NOTE:",
        "--",
    }

    first_skipped = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Rileva nome pagina: testo tutto maiuscolo (min 2 char, non URL, non label)
        is_page_name = (
            text == text.upper()
            and len(text) >= 2
            and not text.startswith("URL")
            and not text.startswith("HTTP")
            and not text.endswith(":")
            and not re.match(r"^[A-Z]{1,3}:\\", text)
            and text not in SKIP_NAMES
            # non è solo simboli o numeri
            and re.search(r"[A-Z]", text)
        )

        if is_page_name:
            # Salta il primo nome-pagina (titolo globale del documento)
            if not first_skipped:
                first_skipped = True
                continue
            # Salva pagina precedente
            if current_page and current_page not in SKIP_NAMES:
                pages.append({
                    "page": current_page,
                    "url": current_url or _slugify(current_page),
                    "content": " ".join(current_lines).strip(),
                })
            current_page = text
            current_url = ""
            current_lines = []

        elif text.upper().startswith("URL:") or text.upper().startswith("URL :"):
            raw_url = re.sub(r"(?i)^url\s*:", "", text).strip()
            current_url = raw_url if raw_url else _slugify(current_page or "")

        else:
            current_lines.append(text)

    # Ultima pagina
    if current_page and current_page not in SKIP_NAMES:
        pages.append({
            "page": current_page,
            "url": current_url or _slugify(current_page),
            "content": " ".join(current_lines).strip(),
        })

    return pages


# ══════════════════════════════════════════════
#  MODELLI
# ══════════════════════════════════════════════

class PageItem(BaseModel):
    page: str
    url: str
    content: str

class MetaItem(BaseModel):
    page: str
    url: str
    title: str
    description: str

class ExportRequest(BaseModel):
    pages: list[MetaItem]


# ══════════════════════════════════════════════
#  ROUTE
# ══════════════════════════════════════════════

@router.post("/parse")
async def parse_document(
    file: UploadFile = File(...),
    _user=Depends(get_current_user),
):
    """Riceve .docx, estrae le sezioni-pagina e suggerisce gli URL."""
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Carica un file .docx valido.")

    raw = await file.read()
    try:
        doc = Document(io.BytesIO(raw))
    except Exception:
        raise HTTPException(status_code=422, detail="Impossibile leggere il file .docx.")

    pages = _parse_sections(doc)
    if not pages:
        raise HTTPException(status_code=422, detail="Nessuna sezione trovata nel documento. Verifica il formato.")

    return pages


@router.post("/generate")
async def generate_meta(
    pages: list[PageItem],
    _user=Depends(get_current_user),
    x_openai_key: Optional[str] = Header(default=None),
):
    """Genera meta title e description per ogni pagina tramite GPT-4o."""
    if not pages:
        raise HTTPException(status_code=400, detail="Nessuna pagina ricevuta.")

    client = _openai_client(x_openai_key)

    SYSTEM = (
        "Sei un SEO specialist esperto. Genera meta title e meta description ottimizzati "
        "per il posizionamento organico.\n\n"
        "Regole meta title:\n"
        "* Lunghezza ideale: 50-60 caratteri\n"
        "* Includi la keyword principale della pagina\n"
        "* Può includere il brand name alla fine separato da \" | \" se c'è spazio\n"
        "* Deve essere descrittivo e invitare al click\n"
        "* Non usare clickbait\n\n"
        "Regole meta description:\n"
        "* Lunghezza ideale: 140-160 caratteri\n"
        "* Deve riassumere il contenuto della pagina\n"
        "* Includi una call to action implicita\n"
        "* Usa la keyword principale naturalmente\n"
        "* Non ripetere il title\n\n"
        "Rispondi SOLO con JSON valido, nessun testo fuori: "
        "{\"title\": \"...\", \"description\": \"...\"}"
    )

    results: list[dict] = []
    for p in pages:
        user_msg = f"Pagina: {p.page}\nURL: {p.url}\nContenuto: {p.content[:1500]}"
        try:
            resp = await client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": SYSTEM},
                    {"role": "user",   "content": user_msg},
                ],
                temperature=0.3,
                max_tokens=300,
            )
            raw_json = resp.choices[0].message.content or "{}"
            # Pulisci eventuale markdown fence
            raw_json = re.sub(r"^```[a-z]*\n?|\n?```$", "", raw_json.strip())
            data = json.loads(raw_json)
            results.append({
                "page":        p.page,
                "url":         p.url,
                "title":       data.get("title", ""),
                "description": data.get("description", ""),
            })
        except Exception as e:
            results.append({
                "page":        p.page,
                "url":         p.url,
                "title":       "",
                "description": "",
                "error":       str(e),
            })

    return results


@router.post("/regenerate")
async def regenerate_single(
    page: PageItem,
    _user=Depends(get_current_user),
    x_openai_key: Optional[str] = Header(default=None),
):
    """Rigenera i meta per una singola pagina."""
    result = await generate_meta([page], _user=_user, x_openai_key=x_openai_key)
    return result[0] if result else {}


@router.post("/export")
async def export_docx(
    body: ExportRequest,
    _user=Depends(get_current_user),
):
    """Genera e restituisce il .docx con i meta tag."""
    pages = body.pages
    if not pages:
        raise HTTPException(status_code=400, detail="Nessuna pagina da esportare.")

    doc = Document()

    # Stile titolo documento
    title_para = doc.add_heading(f"Meta Tag — {date.today().strftime('%d/%m/%Y')}", level=0)
    title_para.runs[0].font.size = Pt(18)

    for p in pages:
        doc.add_heading(p.page, level=1)

        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"

        # Intestazione
        hdr = table.rows[0].cells
        hdr[0].text = "Campo"
        hdr[1].text = "Valore"
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True

        rows_data = [
            ("URL",              p.url),
            ("Meta Title",       p.title),
            ("Meta Description", p.description),
        ]
        for i, (label, value) in enumerate(rows_data, start=1):
            row = table.rows[i].cells
            row[0].text = label
            row[1].text = value

        # Imposta larghezza colonne
        for row in table.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(5.0)

        doc.add_paragraph()  # spazio tra pagine

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"meta-tag-{date.today().isoformat()}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
